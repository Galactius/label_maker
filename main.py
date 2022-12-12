# Python script developed by Daniel Perez, Villanova Vatican Internship Program 2022 - Assignment: Augustinians
import sys
import os
from docx import Document                    # python-docx, enables R/W access for Word Docs.
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert                # docx2pdf, enables conversion between docx and pdf
from datetime import date                   # datetime, used to make use of dates
import tkinter as tk                        # tkinter, Python GUI framework
from tkinter import simpledialog
from tkinter import messagebox


class MainApplication(tk.Frame):
    def __init__(self, parent, *args, **kwargs):
        # Setup inheritance for tkinter and remove initialized window
        tk.Frame.__init__(self, parent, *args, **kwargs)
        self.parent = parent
        self.parent.title('My Window')
        root.withdraw()

        # Get current date/time
        date_str = date.today()

        # Setup Word Document and Paragraph for labels
        document = Document()
        label = document.add_paragraph()

        labels_list = []

        # Setup styling for thicker label
        font_styles = document.styles
        font_charstyle = font_styles.add_style('thick_label', WD_STYLE_TYPE.CHARACTER)
        font_object = font_charstyle.font
        font_object.size = Pt(25)
        font_object.bold = True
        font_object.name = 'Times New Roman'

        # Setup styling for date label
        font_styles2 = document.styles
        font_charstyle2 = font_styles2.add_style('date', WD_STYLE_TYPE.CHARACTER)
        font_object2 = font_charstyle2.font
        font_object2.size = Pt(10)
        font_object2.bold = False
        font_object2.name = 'Times New Roman'

        for x in range(3):
            # Create Dialog Prompts
            dlg_name = simpledialog.askstring('Name Label ' + str(x + 1) + ' of 3', 'Enter Guest Name: (Press enter to skip)', parent=parent)
            dlg_room_no = simpledialog.askinteger('Room No. Label ' + str(x + 1) + ' of 3', 'Enter Guest Room Number: (Enter 0 to skip)', parent=parent)
            # dlg_date = simpledialog.askstring('Date Label ' + str(x + 1) + ' of 3', 'Enter dates for stay MM/DD/ - MM/DD: (Press Enter to skip)', parent=parent)

            if(dlg_name is None or dlg_room_no is None):
                sys.exit("Operation Cancelled (user pressed cancel)")

            string = dlg_name + "  " + str(dlg_room_no)
            #dur_str = str(dlg_date)
            labels_list.append(string)

            # Create confirmation message and end tkinter root
            confirm = messagebox.showinfo('Confirmation', 'Success! Generating label: ' + string)
        root.destroy()

        # Remove "skipped" labels
        for i in labels_list[:]:
            if (i == "  0"):
                labels_list.remove(i)

        # Print data to word file
        for x in range(len(labels_list)):
            cur_str = labels_list[x]

            # Add text to docx and save
            label.add_run(cur_str, style='thick_label')
            # label.add_run(dur_str, style='date')
            label.add_run("\n")
            label.add_run("\n")
            label.add_run("\n")
            label.add_run("\n")
        # Name new label document and save to file
        cur_document = "labels - " + str(date_str) + ".docx"
        document.save(cur_document)

        # Covert document to pdf for printing (to physical printer) then delete printed pdf.
        # needed to resolve XPS error when printing .docx files
        convert(cur_document, "labels.pdf")
        os.system("lp labels.pdf")
        os.remove("labels.pdf")


if __name__ == '__main__':
    root = tk.Tk()
    MainApplication(root)
    root.mainloop()
